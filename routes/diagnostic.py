from __future__ import annotations

import json
import os
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from flask import (
    Blueprint,
    abort,
    after_this_request,
    redirect,
    render_template,
    request,
    send_file,
    url_for,
)
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

from services.diagnostics import (
    get_active_input_pack,
    get_diagnostic_run_by_token,
    get_input_pack_attachments,
    get_latest_input_pack,
    save_client_input_pack,
    save_diagnostic_attachment,
    upsert_active_input_pack,
)
from services.site_links import get_site_links


diagnostic_bp = Blueprint(
    "diagnostic",
    __name__,
    url_prefix="/consulting/diagnostic",
)


BASE_DIR = Path(__file__).resolve().parent.parent
UPLOAD_ROOT = BASE_DIR / "uploads" / "diagnostics"

DOWNLOAD_ROOT = BASE_DIR / "static" / "consulting" / "downloads"

ALLOWED_EXTENSIONS = {
    "xlsx",
    "xls",
    "csv",
    "pdf",
    "doc",
    "docx",
    "txt",
    "png",
    "jpg",
    "jpeg",
}

FIELD_LABELS = {
    "client.company": "Компания",
    "client.contact_name": "Контактное лицо",
    "client.contact_email": "Email",
    "client.process_owner": "Владелец процесса",

    "diagnostic_goal.goal_automation": "Цель: автоматизация",
    "diagnostic_goal.goal_ai_feasibility": "Цель: проверка применимости AI",
    "diagnostic_goal.goal_economics": "Цель: экономика",
    "diagnostic_goal.goal_bottlenecks": "Цель: узкие места",
    "diagnostic_goal.goal_mvp_scope": "Цель: scope MVP",
    "diagnostic_goal.goal_other": "Другая цель",

    "process.process_name": "Название процесса",
    "process.process_description": "Описание процесса",
    "process.main_problem": "Главная проблема",
    "process.request_start": "Как начинается заявка / операция",
    "process.request_channels": "Каналы поступления заявок",
    "process.registration_place": "Где регистрируется заявка",
    "process.roles_description": "Роли в процессе",
    "process.statuses_description": "Статусы процесса",
    "process.sla_description": "SLA / сроки",
    "process.manual_operations": "Ручные операции",
    "process.bottlenecks": "Узкие места",

    "data.excel_log_available": "Есть ли журнал / Excel / выгрузка",
    "data.data_period": "Период данных",
    "data.approx_rows": "Примерное количество строк",
    "data.data_owner": "Владелец данных",
    "data.has_request_id": "Есть ID заявки",
    "data.has_timestamps": "Есть временные метки",
    "data.has_statuses": "Есть статусы",
    "data.has_responsible": "Есть ответственные",
    "data.has_category": "Есть категории",
    "data.has_result": "Есть результат обработки",
    "data.has_free_text": "Есть свободный текст",
    "data.data_quality_issues": "Проблемы качества данных",

    "systems.systems_used": "Используемые системы",
    "systems.systems_description": "Описание систем",

    "integrations.api_available": "Есть API",
    "integrations.exports_available": "Есть выгрузки",
    "integrations.manual_exchange": "Есть ручной обмен",
    "integrations.integration_description": "Описание интеграций",
    "integrations.it_contact": "ИТ-контакт",

    "security.has_personal_data": "Есть персональные данные",
    "security.personal_data_description": "Описание персональных данных",
    "security.can_anonymize": "Можно обезличить",
    "security.nda_required": "Требуется NDA",
    "security.nda_signed": "NDA подписан",
    "security.cloud_allowed": "Облако допустимо",
    "security.security_requirements": "Требования ИБ",
    "security.personal_data_requirements": "Требования к персональным данным",

    "economics.monthly_requests": "Операций в месяц",
    "economics.weekly_requests": "Операций в неделю",
    "economics.employees_involved": "Сотрудников в процессе",
    "economics.avg_processing_time": "Среднее время обработки",
    "economics.monthly_hours": "Часов в месяц",
    "economics.hour_cost": "Стоимость часа",
    "economics.losses_from_errors": "Потери от ошибок",
    "economics.losses_from_delays": "Потери от задержек",
    "economics.expected_effect": "Ожидаемый эффект",

    "contacts.process_contact": "Контакт по процессу",
    "contacts.data_contact": "Контакт по данным",
    "contacts.it_contact": "ИТ-контакт",
    "contacts.security_contact": "Контакт по ИБ",
    "contacts.finance_contact": "Финансовый контакт",

    "client_questions": "Вопросы клиента",

    "confirmation.data_usage_confirmed": "Согласие на использование данных",
    "confirmation.limitations": "Ограничения",
    "confirmation.responsible_person": "Ответственное лицо",
}

BLANK_FORM_SECTIONS = [
    (
        "1. Общая информация",
        [
            "Компания",
            "Контактное лицо",
            "Email контактного лица",
            "Владелец процесса",
        ],
    ),
    (
        "2. Цель диагностики",
        [
            "Понять, можно ли автоматизировать процесс",
            "Оценить целесообразность AI-решения",
            "Проверить экономический эффект",
            "Найти узкие места процесса",
            "Подготовить MVP scope",
            "Другое",
        ],
    ),
    (
        "3. Описание процесса",
        [
            "Название процесса",
            "Как появляется заявка",
            "Краткое описание процесса",
            "Основная проблема процесса",
            "Каналы поступления заявок",
            "Где регистрируется заявка",
            "SLA / сроки обработки",
            "Роли в процессе",
            "Статусы заявки",
            "Ручные операции",
            "Узкие места и потери",
        ],
    ),
    (
        "4. Данные и Excel-журнал",
        [
            "Есть ли Excel-журнал заявок",
            "Период данных",
            "Примерное количество строк",
            "Кто владелец данных",
            "Есть ID заявки",
            "Есть временные метки",
            "Есть статусы",
            "Есть ответственный",
            "Есть категория заявки",
            "Есть результат обработки",
            "Есть свободный текст",
            "Известные проблемы качества данных",
        ],
    ),
    (
        "5. Используемые системы",
        [
            "Excel",
            "1C",
            "CRM",
            "Email",
            "Telegram",
            "Helpdesk / Service Desk",
            "Внутренняя система",
            "Другое",
            "Описание систем",
        ],
    ),
    (
        "6. Интеграции и API",
        [
            "Есть ли API",
            "Есть ли выгрузки Excel/CSV",
            "Есть ли ручной обмен файлами",
            "Контакт ИТ",
            "Описание интеграций",
        ],
    ),
    (
        "7. Информационная безопасность и ПДн",
        [
            "Есть ли персональные данные",
            "Можно ли обезличить данные",
            "Требуется ли NDA",
            "NDA подписан",
            "Разрешена облачная обработка",
            "Какие ПДн могут встречаться",
            "Требования ИБ",
            "Требования по ПДн",
        ],
    ),
    (
        "8. Экономика процесса",
        [
            "Заявок в месяц",
            "Заявок в неделю",
            "Сотрудников участвует",
            "Среднее время обработки заявки",
            "Часов в месяц на процесс",
            "Стоимость часа сотрудника / команды",
            "Потери от ошибок",
            "Потери от задержек / просрочек",
            "Какой эффект от автоматизации считается значимым",
        ],
    ),
    (
        "9. Материалы и файлы",
        [
            "Excel-журнал заявок",
            "Регламент",
            "Описание систем",
            "Описание интеграций/API",
            "Политики ИБ",
            "NDA",
            "Метрики процесса",
        ],
    ),
    (
        "10. Контакты для уточнений",
        [
            "Контакт по процессу",
            "Контакт по данным",
            "Контакт по ИБ / ПДн",
            "Финансовый / экономический контакт",
        ],
    ),
    (
        "11. Открытые вопросы",
        [
            "Вопросы, комментарии, ограничения",
        ],
    ),
    (
        "12. Подтверждение",
        [
            "Данные можно использовать для диагностики",
            "Ответственный со стороны клиента",
            "Ограничения на использование данных",
        ],
    ),
]

def _is_allowed_file(filename: str) -> bool:
    if "." not in filename:
        return False

    extension = filename.rsplit(".", 1)[1].lower()
    return extension in ALLOWED_EXTENSIONS


def _detect_file_type(filename: str) -> str:
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if extension in {"xlsx", "xls", "csv"}:
        return "DATA_FILE"

    if extension in {"doc", "docx", "pdf"}:
        return "DOCUMENT"

    if extension in {"png", "jpg", "jpeg"}:
        return "IMAGE"

    return "OTHER"


def _extract_input_pack_id(saved_input_pack: Any) -> int:
    if isinstance(saved_input_pack, int):
        return saved_input_pack

    if isinstance(saved_input_pack, dict):
        return int(saved_input_pack["id"])

    if hasattr(saved_input_pack, "keys") and "id" in saved_input_pack.keys():
        return int(saved_input_pack["id"])

    return int(saved_input_pack)


def _save_uploaded_files(
    files: list[FileStorage],
    diagnostic_run_id: int,
    input_pack_id: int,
) -> None:
    upload_dir = UPLOAD_ROOT / str(diagnostic_run_id)
    upload_dir.mkdir(parents=True, exist_ok=True)

    for uploaded_file in files:
        if not uploaded_file or not uploaded_file.filename:
            continue

        original_filename = uploaded_file.filename

        if not _is_allowed_file(original_filename):
            continue

        safe_name = secure_filename(original_filename)

        if not safe_name:
            continue

        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
        stored_filename = f"{timestamp}_{safe_name}"
        file_path = upload_dir / stored_filename

        uploaded_file.save(file_path)

        save_diagnostic_attachment(
            diagnostic_run_id=diagnostic_run_id,
            input_pack_id=input_pack_id,
            file_type=_detect_file_type(original_filename),
            original_filename=original_filename,
            stored_filename=stored_filename,
            file_path=str(file_path),
        )


def _build_input_pack_payload() -> dict[str, Any]:
    return {
        "client": {
            "company": request.form.get("company"),
            "contact_name": request.form.get("contact_name"),
            "contact_email": request.form.get("contact_email"),
            "process_owner": request.form.get("process_owner"),
        },
        "diagnostic_goal": {
            "goal_automation": request.form.get("goal_automation"),
            "goal_ai_feasibility": request.form.get("goal_ai_feasibility"),
            "goal_economics": request.form.get("goal_economics"),
            "goal_bottlenecks": request.form.get("goal_bottlenecks"),
            "goal_mvp_scope": request.form.get("goal_mvp_scope"),
            "goal_other": request.form.get("goal_other"),
        },
        "process": {
            "process_name": request.form.get("process_name"),
            "process_description": request.form.get("process_description"),
            "main_problem": request.form.get("main_problem"),
            "request_start": request.form.get("request_start"),
            "request_channels": request.form.getlist("request_channels"),
            "registration_place": request.form.get("registration_place"),
            "roles_description": request.form.get("roles_description"),
            "statuses_description": request.form.get("statuses_description"),
            "sla_description": request.form.get("sla_description"),
            "manual_operations": request.form.get("manual_operations"),
            "bottlenecks": request.form.get("bottlenecks"),
        },
        "data": {
            "excel_log_available": request.form.get("excel_log_available"),
            "data_period": request.form.get("data_period"),
            "approx_rows": request.form.get("approx_rows"),
            "data_owner": request.form.get("data_owner"),
            "has_request_id": request.form.get("has_request_id"),
            "has_timestamps": request.form.get("has_timestamps"),
            "has_statuses": request.form.get("has_statuses"),
            "has_responsible": request.form.get("has_responsible"),
            "has_category": request.form.get("has_category"),
            "has_result": request.form.get("has_result"),
            "has_free_text": request.form.get("has_free_text"),
            "data_quality_issues": request.form.get("data_quality_issues"),
        },
        "systems": {
            "systems_used": request.form.getlist("systems_used"),
            "systems_description": request.form.get("systems_description"),
        },
        "integrations": {
            "api_available": request.form.get("api_available"),
            "exports_available": request.form.get("exports_available"),
            "manual_exchange": request.form.get("manual_exchange"),
            "integration_description": request.form.get("integration_description"),
            "it_contact": request.form.get("it_contact"),
        },
        "security": {
            "has_personal_data": request.form.get("has_personal_data"),
            "personal_data_description": request.form.get("personal_data_description"),
            "can_anonymize": request.form.get("can_anonymize"),
            "nda_required": request.form.get("nda_required"),
            "nda_signed": request.form.get("nda_signed"),
            "cloud_allowed": request.form.get("cloud_allowed"),
            "security_requirements": request.form.get("security_requirements"),
            "personal_data_requirements": request.form.get("personal_data_requirements"),
        },
        "economics": {
            "monthly_requests": request.form.get("monthly_requests"),
            "weekly_requests": request.form.get("weekly_requests"),
            "employees_involved": request.form.get("employees_involved"),
            "avg_processing_time": request.form.get("avg_processing_time"),
            "monthly_hours": request.form.get("monthly_hours"),
            "hour_cost": request.form.get("hour_cost"),
            "losses_from_errors": request.form.get("losses_from_errors"),
            "losses_from_delays": request.form.get("losses_from_delays"),
            "expected_effect": request.form.get("expected_effect"),
        },
        "contacts": {
            "process_contact": request.form.get("process_contact"),
            "data_contact": request.form.get("data_contact"),
            "it_contact": request.form.get("it_contact"),
            "security_contact": request.form.get("security_contact"),
            "finance_contact": request.form.get("finance_contact"),
        },
        "client_questions": request.form.get("client_questions"),
        "confirmation": {
            "data_usage_confirmed": request.form.get("data_usage_confirmed"),
            "limitations": request.form.get("limitations"),
            "responsible_person": request.form.get("responsible_person"),
        },
    }


def _load_raw_payload(input_pack: Any) -> dict[str, Any]:
    if input_pack is None:
        return {}

    raw_payload = input_pack["raw_payload"]

    if not raw_payload:
        return {}

    if isinstance(raw_payload, dict):
        return raw_payload

    try:
        payload = json.loads(raw_payload)
    except Exception:
        return {}

    if isinstance(payload, dict):
        return payload

    return {}


def _flatten_payload(
    payload: dict[str, Any],
    prefix: str = "",
) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []

    for key, value in payload.items():
        full_key = f"{prefix}.{key}" if prefix else key

        if isinstance(value, dict):
            rows.extend(_flatten_payload(value, full_key))
            continue

        label = FIELD_LABELS.get(full_key, full_key)

        if isinstance(value, list):
            clean_value = ", ".join(str(item) for item in value if item)
        elif value is None or value == "":
            clean_value = "не указано"
        else:
            clean_value = str(value)

        rows.append((label, clean_value))

    return rows


def _docx_paragraph(text: str, bold: bool = False) -> str:
    safe_text = escape(text)

    if bold:
        return (
            "<w:p>"
            "<w:r>"
            "<w:rPr><w:b/></w:rPr>"
            f"<w:t>{safe_text}</w:t>"
            "</w:r>"
            "</w:p>"
        )

    return (
        "<w:p>"
        "<w:r>"
        f"<w:t>{safe_text}</w:t>"
        "</w:r>"
        "</w:p>"
    )

def _build_blank_docx(
    output_path: Path,
    diagnostic_run: Any,
) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Inches, Pt, RGBColor
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Для генерации DOCX установите зависимость: pip install python-docx"
        ) from exc

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Diagnostic Input Pack Form")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AIha Consulting")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 95, 120)

    document.add_paragraph()

    info_table = document.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = "Table Grid"

    info_rows = [
        ("Diagnostic ID", str(diagnostic_run["id"])),
        ("Компания", diagnostic_run["company"] or "не указано"),
        ("Контакт", diagnostic_run["contact_name"] or "не указано"),
        ("Email", diagnostic_run["contact_email"] or "не указано"),
    ]

    for row_index, (label, value) in enumerate(info_rows):
        cells = info_table.rows[row_index].cells
        cells[0].text = label
        cells[1].text = value

        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)

        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(17, 24, 39)

    document.add_paragraph()

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = intro.add_run(
        "Заполните форму внутри компании. Её можно использовать для согласования "
        "с владельцем процесса, ИТ, ИБ и финансовым блоком. После подготовки "
        "ответы можно перенести в онлайн-форму или отправить документ специалисту AIha Consulting."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(57, 65, 80)

    document.add_paragraph()

    for section_title, fields in BLANK_FORM_SECTIONS:
        heading = document.add_paragraph()
        heading_run = heading.add_run(section_title)
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(17, 24, 39)

        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header_cells = table.rows[0].cells
        header_cells[0].text = "Поле"
        header_cells[1].text = "Ответ / комментарий"

        for cell in header_cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        for field in fields:
            row = table.add_row()
            row.cells[0].text = field
            row.cells[1].text = ""

            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)

        for row in table.rows:
            row.height = Pt(26)

        document.add_paragraph()

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("AIha Consulting · Diagnostic Input Pack")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(102, 112, 133)

    document.save(output_path)

def _build_submitted_docx(
    output_path: Path,
    diagnostic_run: Any,
    payload: dict[str, Any],
) -> None:
    rows = _flatten_payload(payload)

    paragraphs = [
        _docx_paragraph("Diagnostic Input Pack — AIha Consulting", bold=True),
        _docx_paragraph(f"Diagnostic ID: {diagnostic_run['id']}"),
        _docx_paragraph(f"Компания: {diagnostic_run['company'] or 'не указано'}"),
        _docx_paragraph(f"Контакт: {diagnostic_run['contact_name'] or 'не указано'}"),
        _docx_paragraph(f"Email: {diagnostic_run['contact_email'] or 'не указано'}"),
        _docx_paragraph(f"Дата выгрузки: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC"),
        _docx_paragraph(""),
        _docx_paragraph("Заполненные данные", bold=True),
    ]

    for label, value in rows:
        paragraphs.append(_docx_paragraph(label, bold=True))
        paragraphs.append(_docx_paragraph(value))
        paragraphs.append(_docx_paragraph(""))

    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document
    xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
    <w:body>
        {''.join(paragraphs)}
        <w:sectPr>
            <w:pgSz w:w="11906" w:h="16838"/>
            <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>
        </w:sectPr>
    </w:body>
</w:document>
"""

    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
"""

    rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship
        Id="rId1"
        Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
        Target="word/document.xml"/>
</Relationships>
"""

    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", content_types_xml)
        docx.writestr("_rels/.rels", rels_xml)
        docx.writestr("word/document.xml", document_xml)

def _make_temp_output_path(suffix: str) -> Path:
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)
    return Path(path)

def _find_pdf_font() -> Path | None:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
    ]

    for candidate in candidates:
        if candidate.exists():
            return candidate

    return None

def _build_blank_pdf(
    output_path: Path,
    diagnostic_run: Any,
) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            PageBreak,
        )
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Для генерации PDF установите зависимость: pip install reportlab"
        ) from exc

    font_path = _find_pdf_font()

    if font_path is None:
        raise RuntimeError(
            "Не найден TTF-шрифт для генерации PDF. "
            "На Windows обычно подходит C:/Windows/Fonts/arial.ttf."
        )

    pdfmetrics.registerFont(TTFont("AIhaFont", str(font_path)))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="Diagnostic Input Pack Form — AIha Consulting",
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "AIhaTitle",
        parent=styles["Title"],
        fontName="AIhaFont",
        fontSize=22,
        leading=26,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#111827"),
        spaceAfter=8,
    )

    subtitle_style = ParagraphStyle(
        "AIhaSubtitle",
        parent=styles["Normal"],
        fontName="AIhaFont",
        fontSize=11,
        leading=15,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#505f78"),
        spaceAfter=18,
    )

    section_style = ParagraphStyle(
        "AIhaSection",
        parent=styles["Heading2"],
        fontName="AIhaFont",
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#111827"),
        spaceBefore=18,
        spaceAfter=10,
    )

    normal_style = ParagraphStyle(
        "AIhaNormal",
        parent=styles["Normal"],
        fontName="AIhaFont",
        fontSize=9,
        leading=13,
        textColor=colors.HexColor("#394150"),
        alignment=TA_LEFT,
    )

    small_style = ParagraphStyle(
        "AIhaSmall",
        parent=styles["Normal"],
        fontName="AIhaFont",
        fontSize=8,
        leading=11,
        textColor=colors.HexColor("#667085"),
        alignment=TA_CENTER,
    )

    story = []

    story.append(Paragraph("Diagnostic Input Pack Form", title_style))
    story.append(Paragraph("AIha Consulting", subtitle_style))

    info_table = Table(
        [
            ["Diagnostic ID", str(diagnostic_run["id"])],
            ["Компания", diagnostic_run["company"] or "не указано"],
            ["Контакт", diagnostic_run["contact_name"] or "не указано"],
            ["Email", diagnostic_run["contact_email"] or "не указано"],
        ],
        colWidths=[45 * mm, 115 * mm],
    )

    info_table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "AIhaFont"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef2f7")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#111827")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d5dd")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )

    story.append(info_table)
    story.append(Spacer(1, 12))

    story.append(
        Paragraph(
            "Заполните форму внутри компании. Её можно использовать для согласования "
            "с владельцем процесса, ИТ, ИБ и финансовым блоком. После подготовки "
            "ответы можно перенести в онлайн-форму или отправить документ специалисту AIha Consulting.",
            normal_style,
        )
    )

    story.append(Spacer(1, 12))

    for section_index, (section_title, fields) in enumerate(BLANK_FORM_SECTIONS, start=1):
        story.append(Paragraph(section_title, section_style))

        table_data = [
            [
                Paragraph("Поле", normal_style),
                Paragraph("Ответ / комментарий", normal_style),
            ]
        ]

        for field in fields:
            table_data.append(
                [
                    Paragraph(field, normal_style),
                    Paragraph(" ", normal_style),
                ]
            )

        section_table = Table(
            table_data,
            colWidths=[58 * mm, 102 * mm],
            repeatRows=1,
        )

        section_table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "AIhaFont"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#111827")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d0d5dd")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 7),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                ]
            )
        )

        story.append(section_table)
        story.append(Spacer(1, 10))

        if section_index in {4, 8}:
            story.append(PageBreak())

    story.append(Spacer(1, 16))
    story.append(Paragraph("AIha Consulting · Diagnostic Input Pack", small_style))

    doc.build(story)


def _build_submitted_pdf(
    output_path: Path,
    diagnostic_run: Any,
    payload: dict[str, Any],
) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Для генерации PDF установите зависимость: pip install reportlab"
        ) from exc

    font_path = _find_pdf_font()

    if font_path is None:
        raise RuntimeError(
            "Не найден TTF-шрифт для генерации PDF. "
            "На Windows обычно подходит C:/Windows/Fonts/arial.ttf."
        )

    pdfmetrics.registerFont(TTFont("AIhaFont", str(font_path)))

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    left = 48
    top = height - 52
    line_height = 16
    y = top

    def draw_line(text: str, bold: bool = False) -> None:
        nonlocal y

        if y < 64:
            c.showPage()
            c.setFont("AIhaFont", 10)
            y = top

        c.setFont("AIhaFont", 12 if bold else 10)

        max_chars = 95
        text = text or ""

        chunks = [
            text[index : index + max_chars]
            for index in range(0, len(text), max_chars)
        ] or [""]

        for chunk in chunks:
            if y < 64:
                c.showPage()
                c.setFont("AIhaFont", 10)
                y = top

            c.drawString(left, y, chunk)
            y -= line_height

    draw_line("Diagnostic Input Pack — AIha Consulting", bold=True)
    draw_line(f"Diagnostic ID: {diagnostic_run['id']}")
    draw_line(f"Компания: {diagnostic_run['company'] or 'не указано'}")
    draw_line(f"Контакт: {diagnostic_run['contact_name'] or 'не указано'}")
    draw_line(f"Email: {diagnostic_run['contact_email'] or 'не указано'}")
    draw_line(f"Дата выгрузки: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC")
    draw_line("")
    draw_line("Заполненные данные", bold=True)
    draw_line("")

    rows = _flatten_payload(payload)

    for label, value in rows:
        draw_line(label, bold=True)
        draw_line(value)
        draw_line("")

    c.save()


@diagnostic_bp.route("/input-pack/<token>", methods=["GET", "POST"])
def input_pack(token: str):
    diagnostic_run = get_diagnostic_run_by_token(token)

    if diagnostic_run is None:
        return render_template(
            "consulting/diagnostic_input_pack_invalid.html",
            site_links=get_site_links(),
        ), 404

    if request.method == "POST":
        payload = _build_input_pack_payload()

        # Жёсткая защита, чтобы базовый AI Audit Brief не сохранился как другой тип brief
        payload["brief_type"] = "diagnostic_input_pack"
        payload["source"] = payload.get("source") or "web_form"

        saved_input_pack = upsert_active_input_pack(
            diagnostic_run_id=diagnostic_run["id"],
            brief_type="diagnostic_input_pack",
            payload=payload,
            source=payload["source"],
        )

        input_pack_id = _extract_input_pack_id(saved_input_pack)

        uploaded_files = request.files.getlist("attachments")
        _save_uploaded_files(
            files=uploaded_files,
            diagnostic_run_id=diagnostic_run["id"],
            input_pack_id=input_pack_id,
        )

        return redirect(
            url_for("diagnostic.input_pack_submitted", token=token)
        )

    active_input_pack = get_active_input_pack(
        diagnostic_run_id=diagnostic_run["id"],
        brief_type="diagnostic_input_pack",
    )

    form_data = _load_raw_payload(active_input_pack)

    existing_attachments = []
    if active_input_pack is not None:
        existing_attachments = get_input_pack_attachments(
            int(active_input_pack["id"])
        )

    return render_template(
        "consulting/diagnostic_input_pack.html",
        site_links=get_site_links(),
        diagnostic_run=diagnostic_run,
        diagnostic=diagnostic_run,
        token=token,
        form_data=form_data,
        edit_mode=active_input_pack is not None,
        input_pack_id=active_input_pack["id"] if active_input_pack else None,
        existing_attachments=existing_attachments,
        download_docx_url=url_for(
            "diagnostic.download_input_pack_template",
            token=token,
            file_format="docx",
        ),
        download_pdf_url=url_for(
            "diagnostic.download_input_pack_template",
            token=token,
            file_format="pdf",
        ),
    )

@diagnostic_bp.route("/input-pack/<token>/download/<file_format>")
def download_input_pack_template(token: str, file_format: str):
    diagnostic_run = get_diagnostic_run_by_token(token)

    if diagnostic_run is None:
        abort(404)

    file_format = file_format.lower().strip()

    if file_format not in {"docx", "pdf"}:
        abort(404)

    suffix = f".{file_format}"
    output_path = _make_temp_output_path(suffix)

    try:
        if file_format == "docx":
            _build_blank_docx(
                output_path=output_path,
                diagnostic_run=diagnostic_run,
            )
            download_name = "AIha_Diagnostic_Input_Pack_Form_v1.docx"
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            _build_blank_pdf(
                output_path=output_path,
                diagnostic_run=diagnostic_run,
            )
            download_name = "AIha_Diagnostic_Input_Pack_Form_v1.pdf"
            mimetype = "application/pdf"

    except RuntimeError as exc:
        output_path.unlink(missing_ok=True)
        return str(exc), 500

    response = send_file(
        output_path,
        as_attachment=True,
        download_name=download_name,
        mimetype=mimetype,
    )

    response.call_on_close(lambda: output_path.unlink(missing_ok=True))

    return response


@diagnostic_bp.route("/input-pack/<token>/submitted")
def input_pack_submitted(token: str):
    diagnostic_run = get_diagnostic_run_by_token(token)

    if diagnostic_run is None:
        abort(404)

    return render_template(
        "consulting/diagnostic_input_pack_submitted.html",
        site_links=get_site_links(),
        diagnostic_run=diagnostic_run,
        diagnostic=diagnostic_run,
        token=token,
        download_docx_url=url_for(
            "diagnostic.download_submitted_input_pack",
            token=token,
            file_format="docx",
        ),
        download_pdf_url=url_for(
            "diagnostic.download_submitted_input_pack",
            token=token,
            file_format="pdf",
        ),
    )


@diagnostic_bp.route("/input-pack/<token>/submitted/download/<file_format>")
def download_submitted_input_pack(token: str, file_format: str):
    diagnostic_run = get_diagnostic_run_by_token(token)

    if diagnostic_run is None:
        abort(404)

    file_format = file_format.lower().strip()

    if file_format not in {"docx", "pdf"}:
        abort(404)

    input_pack = get_active_input_pack(
        diagnostic_run["id"],
        "diagnostic_input_pack",
    )

    if input_pack is None:
        return "Заполненная форма не найдена", 404

    payload = _load_raw_payload(input_pack)

    suffix = f".{file_format}"
    output_path = _make_temp_output_path(suffix)

    try:
        if file_format == "docx":
            _build_submitted_docx(
                output_path=output_path,
                diagnostic_run=diagnostic_run,
                payload=payload,
            )
            download_name = f"AIha_Diagnostic_Input_Pack_Submitted_{diagnostic_run['id']}.docx"
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            _build_submitted_pdf(
                output_path=output_path,
                diagnostic_run=diagnostic_run,
                payload=payload,
            )
            download_name = f"AIha_Diagnostic_Input_Pack_Submitted_{diagnostic_run['id']}.pdf"
            mimetype = "application/pdf"

    except RuntimeError as exc:
        output_path.unlink(missing_ok=True)
        return str(exc), 500

    response = send_file(
        output_path,
        as_attachment=True,
        download_name=download_name,
        mimetype=mimetype,
    )

    response.call_on_close(lambda: output_path.unlink(missing_ok=True))

    return response

INDUSTRIAL_AI_VALUE_LABELS = {
    "failure_prediction": "Прогнозирование поломок оборудования",
    "predictive_maintenance": "Предиктивное обслуживание",
    "downtime_analysis": "Анализ простоев",
    "schedule_optimization": "Оптимизация производственного графика",
    "production_plan_fact": "План-факт производства",
    "quality_control": "Контроль качества / анализ брака",
    "supply_chain": "Анализ цепочки поставок",
    "inventory": "Управление запасами",
    "energy": "Энергопотребление",
    "other": "Другое",

    "reduce_downtime": "Снизить простои",
    "reduce_repairs": "Снизить аварийные ремонты",
    "reduce_defects": "Сократить брак",
    "increase_oee": "Повысить OEE",
    "improve_otd": "Улучшить выполнение заказов в срок",
    "reduce_inventory": "Снизить запасы",
    "reduce_stockouts": "Сократить дефициты",
    "improve_planning": "Повысить точность планирования",
    "reduce_energy_cost": "Снизить энергозатраты",
    "reduce_manual_work": "Снизить ручной труд",

    "erp": "ERP",
    "1c": "1C",
    "mes": "MES",
    "scada": "SCADA",
    "wms": "WMS",
    "excel": "Excel",
    "repair_logs": "Журналы ремонтов",
    "shift_logs": "Журналы смен",
    "downtime_logs": "Журналы простоев",
    "quality_logs": "Журналы качества / ОТК",
    "sensor_data": "Данные датчиков",
    "supply_data": "Данные закупок / поставщиков",

    "duplicates": "Дубли",
    "missing_values": "Пропуски",
    "manual_input": "Ручной ввод",
    "inconsistent_statuses": "Разные статусы",
    "no_single_id": "Нет единого ID",
    "no_system_links": "Нет связки между системами",
    "no_downtime_reasons": "Нет причин простоев",
    "no_defect_reasons": "Нет причин брака",
    "paper_logs": "Данные в бумажных журналах",
    "unknown": "Требует уточнения",

    "event": "По событию",
    "daily": "Ежедневно",
    "shift": "Посменно",
    "hourly": "Ежечасно",
    "minute": "Поминутно",
    "realtime": "В реальном времени",
    "mixed": "Смешанная частота",

    "yes": "Да",
    "no": "Нет",
    "partial": "Частично",

    "required": "Требуется",
    "signed": "Уже подписан",
    "not_required": "Не требуется",

    "anonymized_only": "Только обезличенно",

    "local": "Локально",
    "client_server": "Сервер клиента",
    "aiha_server": "Сервер AIha",
    "cloud": "Облако",
    "hybrid": "Гибрид",

    "employee_personal_data": "ПДн сотрудников",
    "client_personal_data": "ПДн клиентов",
    "commercial_secret": "Коммерческая тайна",
    "cost_price": "Себестоимость",
    "prices": "Цены",
    "suppliers": "Поставщики",
    "recipes": "Рецептуры / технологии",
    "drawings": "Чертежи",
    "production_volume": "Объёмы производства",
    "none": "Нет / не ожидается",

    "downtime": "Простои",
    "repairs": "Аварийные ремонты",
    "defects": "Брак",
    "rework": "Повторная обработка",
    "late_orders": "Просрочки заказов",
    "stockouts": "Дефициты",
    "excess_inventory": "Избыточные запасы",
    "manual_work": "Ручной труд",
    "energy": "Энергозатраты",
    "claims": "Рекламации",

    "2_weeks": "2 недели",
    "4_weeks": "4 недели",
    "6_weeks": "6 недель",
}


INDUSTRIAL_AI_BRIEF_SECTIONS = [
    (
        "1. Тип производственного AI-кейса",
        [
            ("Основной тип кейса", "case_type.primary_case_type"),
            ("Другое / уточнение", "case_type.primary_case_type_other"),
            ("Главные бизнес-цели", "case_type.business_goals"),
            ("Ключевой ожидаемый эффект", "case_type.expected_effect"),
        ],
    ),
    (
        "2. Производственный контекст",
        [
            ("Компания / подразделение", "production_context.company_name"),
            ("Производственная площадка", "production_context.production_site"),
            ("Цех / участок / линия", "production_context.unit_or_line"),
            ("Владелец процесса", "production_context.process_owner"),
            ("Технический контакт", "production_context.technical_contact"),
            ("Контакт по данным / ИТ", "production_context.data_contact"),
            ("Описание текущего процесса", "production_context.current_process_description"),
            ("Где начинается процесс", "production_context.process_start"),
            ("Где заканчивается процесс", "production_context.process_end"),
            ("Что входит в анализ", "production_context.in_scope"),
            ("Что не входит в анализ", "production_context.out_of_scope"),
        ],
    ),
    (
        "3. Объект анализа — оборудование",
        [
            ("Тип оборудования", "analysis_object.equipment.equipment_type"),
            ("Количество единиц оборудования", "analysis_object.equipment.equipment_count"),
            ("Критичное оборудование", "analysis_object.equipment.critical_equipment"),
            ("Режим работы", "analysis_object.equipment.operation_mode"),
            ("Основные отказы / поломки", "analysis_object.equipment.main_failures"),
            ("Средняя длительность простоя", "analysis_object.equipment.average_downtime"),
            ("Стоимость часа простоя", "analysis_object.equipment.downtime_hour_cost"),
        ],
    ),
    (
        "4. Объект анализа — линии, операции, материалы",
        [
            ("Количество линий / участков", "analysis_object.lines_operations.lines_count"),
            ("Основные операции", "analysis_object.lines_operations.main_operations"),
            ("Узкие места", "analysis_object.lines_operations.bottlenecks"),
            ("Частота переналадок", "analysis_object.lines_operations.changeover_frequency"),
            ("Среднее время переналадки", "analysis_object.lines_operations.changeover_duration"),
            ("Тип продукции / SKU", "analysis_object.products_materials.product_types"),
            ("Есть ли партии / batch tracking", "analysis_object.products_materials.batch_tracking"),
            ("Критичные материалы / компоненты", "analysis_object.products_materials.critical_materials"),
            ("Критичные поставщики", "analysis_object.products_materials.critical_suppliers"),
            ("Проблемы с материалами / поставками", "analysis_object.products_materials.material_issues"),
        ],
    ),
    (
        "5. Данные",
        [
            ("Доступные источники данных", "data.data_sources"),
            ("Период исторических данных", "data.data_period"),
            ("Частота данных", "data.data_frequency"),
            ("Идентификаторы в данных", "data.data_identifiers"),
            ("Временные метки", "data.data_timestamps"),
            ("Проблемы качества данных", "data.data_quality_issues"),
        ],
    ),
    (
        "6. Системы и интеграции",
        [
            ("1C / ERP", "systems_integrations.erp_status"),
            ("MES", "systems_integrations.mes_status"),
            ("SCADA", "systems_integrations.scada_status"),
            ("WMS", "systems_integrations.wms_status"),
            ("Excel / CSV выгрузка", "systems_integrations.excel_export_status"),
            ("API", "systems_integrations.api_status"),
            ("Интеграции, критичные для MVP", "systems_integrations.critical_integrations"),
            ("Интеграции, которые можно отложить", "systems_integrations.deferred_integrations"),
        ],
    ),
    (
        "7. ИБ, ПДн и коммерческая тайна",
        [
            ("Чувствительные данные", "security.sensitive_data"),
            ("NDA", "security.nda_status"),
            ("Разрешена облачная обработка", "security.cloud_allowed"),
            ("Можно передавать данные AIha", "security.data_transfer_allowed"),
            ("Предпочтительная обработка", "security.preferred_processing"),
            ("Что нужно обезличить / удалить", "security.anonymization_requirements"),
            ("Дополнительные ограничения ИБ", "security.security_notes"),
        ],
    ),
    (
        "8. Экономика проблемы",
        [
            ("Основная зона потерь", "economics.loss_areas"),
            ("Количество простоев в месяц", "economics.downtime_events_per_month"),
            ("Процент брака", "economics.defect_rate"),
            ("Стоимость брака в месяц", "economics.defect_cost_per_month"),
            ("Стоимость запасов", "economics.inventory_value"),
            ("Ручные трудозатраты", "economics.manual_labor_hours"),
            ("Оценка потерь в месяц", "economics.monthly_loss_estimate"),
            ("Какой эффект считается значимым", "economics.meaningful_effect"),
        ],
    ),
    (
        "9. MVP-рамка",
        [
            ("Какой один кейс берём в первый MVP", "mvp.mvp_case"),
            ("Почему именно этот кейс", "mvp.mvp_reason"),
            ("Участок / линия / оборудование", "mvp.mvp_unit"),
            ("Период данных для MVP", "mvp.mvp_data_period"),
            ("Период теста", "mvp.mvp_test_period"),
            ("Что остаётся под контролем человека", "mvp.human_control"),
            ("Что точно не входит в MVP", "mvp.mvp_exclusions"),
            ("Критерии успеха MVP", "mvp.mvp_success_criteria"),
            ("GO-критерий после MVP", "mvp.mvp_go_criterion"),
            ("NO-GO-критерий после MVP", "mvp.mvp_no_go_criterion"),
        ],
    ),
    (
        "10. Файлы и материалы",
        [
            ("Выгрузки ERP / 1C / MES / SCADA", "attachments.erp_1c_mes_scada_exports"),
            ("Журналы ремонтов", "attachments.repair_logs"),
            ("Журналы простоев", "attachments.downtime_logs"),
            ("Данные по браку / качеству", "attachments.quality_defect_data"),
            ("План-факт производства", "attachments.production_plan_fact"),
            ("Данные по поставкам / материалам", "attachments.supply_material_data"),
            ("Скриншоты систем", "attachments.system_screenshots"),
            ("Описание процесса", "attachments.process_description"),
            ("Регламенты / инструкции", "attachments.regulations_instructions"),
            ("Заполненный DOCX/PDF бланк", "attachments.filled_docx_pdf"),
            ("Дополнительные материалы", "attachments.additional_materials"),
        ],
    ),
    (
        "11. Дополнительно",
        [
            ("Дополнительные комментарии", "additional_comments"),
        ],
    ),
]

def _industrial_ai_payload(payload: dict[str, Any]) -> dict[str, Any]:
    if not isinstance(payload, dict):
        return {}

    industrial_ai = payload.get("industrial_ai")

    if isinstance(industrial_ai, dict):
        return industrial_ai

    return payload


def _industrial_ai_get(payload: dict[str, Any], path: str) -> Any:
    current: Any = _industrial_ai_payload(payload)

    for part in path.split("."):
        if not isinstance(current, dict):
            return None

        current = current.get(part)

    return current


def _industrial_ai_format_value(value: Any) -> str:
    if value is None:
        return "не указано"

    if isinstance(value, list):
        if not value:
            return "не указано"

        return ", ".join(
            INDUSTRIAL_AI_VALUE_LABELS.get(str(item), str(item))
            for item in value
            if str(item).strip()
        ) or "не указано"

    if isinstance(value, bool):
        return "Да" if value else "Нет"

    text = str(value).strip()

    if not text:
        return "не указано"

    return INDUSTRIAL_AI_VALUE_LABELS.get(text, text)


def _industrial_ai_docx_add_paragraph(
    document: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: int | None = None,
) -> Any:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic

    if size is not None:
        from docx.shared import Pt

        run.font.size = Pt(size)

    return paragraph


def _industrial_ai_docx_add_table(document: Any, rows: list[tuple[str, str]]) -> Any:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Поле"
    header_cells[1].text = "Значение"

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    return table


def _industrial_ai_pdf_register_font() -> str:
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError as exc:
        raise RuntimeError(
            "Для генерации PDF установите зависимость: pip install reportlab"
        ) from exc

    font_candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        "/System/Library/Fonts/Arial.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]

    for font_path in font_candidates:
        if Path(font_path).exists():
            font_name = "AIhaIndustrialSans"

            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, font_path))

            return font_name

    return "Helvetica"


def _industrial_ai_pdf_paragraph(text: str, style: Any) -> Any:
    from html import escape as html_escape
    from reportlab.platypus import Paragraph

    safe_text = html_escape(text or "").replace("\n", "<br/>")
    return Paragraph(safe_text, style)


def _build_industrial_ai_blank_docx(
    *,
    output_path: Path,
    diagnostic_run: dict[str, Any],
) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "Для генерации DOCX установите зависимость: pip install python-docx"
        ) from exc

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = document.add_heading("Industrial AI Brief", level=0)
    title.runs[0].font.name = "Arial"

    _industrial_ai_docx_add_paragraph(
        document,
        "Специализированный brief для производственных AI-кейсов. "
        "Заполняется как приложение к Base AI Audit Brief.",
        italic=True,
    )

    _industrial_ai_docx_add_paragraph(
        document,
        f"Diagnostic ID: {diagnostic_run.get('id', '—')}",
    )
    _industrial_ai_docx_add_paragraph(document, "Brief type: industrial_ai")
    _industrial_ai_docx_add_paragraph(document, "Brief version: v1")

    document.add_paragraph()

    _industrial_ai_docx_add_paragraph(
        document,
        "Инструкция: заполните только релевантные поля. "
        "Если информация отсутствует, укажите «не указано» или «требует уточнения».",
        bold=True,
    )

    for section_title, fields in INDUSTRIAL_AI_BRIEF_SECTIONS:
        document.add_heading(section_title, level=1)

        rows = []
        for label, _path in fields:
            rows.append((label, ""))

        _industrial_ai_docx_add_table(document, rows)
        document.add_paragraph()

    document.add_page_break()

    document.add_heading("Вопросы для созвона", level=1)

    questions = [
        "Какой производственный кейс является самым болезненным сейчас?",
        "Где есть данные за последние 3–12 месяцев?",
        "Можно ли начать с выгрузки Excel/CSV без интеграций?",
        "Кто владелец процесса и кто принимает решение?",
        "Какой экономический эффект считается значимым?",
        "Какие данные нельзя передавать наружу?",
        "Можно ли обезличить тестовую выгрузку?",
        "Какие ограничения есть по ИБ и коммерческой тайне?",
        "Какой участок / линия / оборудование подходят для первого MVP?",
        "Какой результат через 4–6 недель будет считаться успехом?",
    ]

    for question in questions:
        document.add_paragraph(question, style="List Number")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _build_industrial_ai_blank_pdf(
    *,
    output_path: Path,
    diagnostic_run: dict[str, Any],
) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            PageBreak,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:
        raise RuntimeError(
            "Для генерации PDF установите зависимость: pip install reportlab"
        ) from exc

    font_name = _industrial_ai_pdf_register_font()

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="AIhaTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=22,
            leading=26,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AIhaHeading",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AIhaBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=9,
            leading=12,
            spaceAfter=6,
        )
    )

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )

    story: list[Any] = []

    story.append(_industrial_ai_pdf_paragraph("Industrial AI Brief", styles["AIhaTitle"]))
    story.append(
        _industrial_ai_pdf_paragraph(
            "Специализированный brief для производственных AI-кейсов. "
            "Заполняется как приложение к Base AI Audit Brief.",
            styles["AIhaBody"],
        )
    )
    story.append(
        _industrial_ai_pdf_paragraph(
            f"Diagnostic ID: {diagnostic_run.get('id', '—')}<br/>"
            "Brief type: industrial_ai<br/>"
            "Brief version: v1",
            styles["AIhaBody"],
        )
    )
    story.append(Spacer(1, 8))

    for section_title, fields in INDUSTRIAL_AI_BRIEF_SECTIONS:
        story.append(_industrial_ai_pdf_paragraph(section_title, styles["AIhaHeading"]))

        data = [
            [
                _industrial_ai_pdf_paragraph("Поле", styles["AIhaBody"]),
                _industrial_ai_pdf_paragraph("Значение", styles["AIhaBody"]),
            ]
        ]

        for label, _path in fields:
            data.append(
                [
                    _industrial_ai_pdf_paragraph(label, styles["AIhaBody"]),
                    _industrial_ai_pdf_paragraph("", styles["AIhaBody"]),
                ]
            )

        table = Table(data, colWidths=[72 * mm, 96 * mm], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEBFF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111111")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8C8D0")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )

        story.append(table)
        story.append(Spacer(1, 10))

    story.append(PageBreak())
    story.append(_industrial_ai_pdf_paragraph("Вопросы для созвона", styles["AIhaHeading"]))

    questions = [
        "1. Какой производственный кейс является самым болезненным сейчас?",
        "2. Где есть данные за последние 3–12 месяцев?",
        "3. Можно ли начать с выгрузки Excel/CSV без интеграций?",
        "4. Кто владелец процесса и кто принимает решение?",
        "5. Какой экономический эффект считается значимым?",
        "6. Какие данные нельзя передавать наружу?",
        "7. Можно ли обезличить тестовую выгрузку?",
        "8. Какие ограничения есть по ИБ и коммерческой тайне?",
        "9. Какой участок / линия / оборудование подходят для первого MVP?",
        "10. Какой результат через 4–6 недель будет считаться успехом?",
    ]

    for question in questions:
        story.append(_industrial_ai_pdf_paragraph(question, styles["AIhaBody"]))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def _build_submitted_industrial_ai_docx(
    *,
    output_path: Path,
    diagnostic_run: dict[str, Any],
    payload: dict[str, Any],
) -> None:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "Для генерации DOCX установите зависимость: pip install python-docx"
        ) from exc

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    document.add_heading("Industrial AI Brief — заполненная форма", level=0)

    _industrial_ai_docx_add_paragraph(
        document,
        f"Diagnostic ID: {diagnostic_run.get('id', '—')}",
    )
    _industrial_ai_docx_add_paragraph(document, "Brief type: industrial_ai")
    _industrial_ai_docx_add_paragraph(
        document,
        f"Brief version: {payload.get('brief_version', 'v1')}",
    )
    _industrial_ai_docx_add_paragraph(
        document,
        f"Submitted at: {payload.get('submitted_at', 'не указано')}",
    )

    document.add_paragraph()

    for section_title, fields in INDUSTRIAL_AI_BRIEF_SECTIONS:
        document.add_heading(section_title, level=1)

        rows = []
        for label, path in fields:
            value = _industrial_ai_format_value(_industrial_ai_get(payload, path))
            rows.append((label, value))

        _industrial_ai_docx_add_table(document, rows)
        document.add_paragraph()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _build_submitted_industrial_ai_pdf(
    *,
    output_path: Path,
    diagnostic_run: dict[str, Any],
    payload: dict[str, Any],
) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError(
            "Для генерации PDF установите зависимость: pip install reportlab"
        ) from exc

    font_name = _industrial_ai_pdf_register_font()

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="AIhaTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=20,
            leading=24,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AIhaHeading",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="AIhaBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=9,
            leading=12,
            spaceAfter=6,
        )
    )

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )

    story: list[Any] = []

    story.append(
        _industrial_ai_pdf_paragraph(
            "Industrial AI Brief — заполненная форма",
            styles["AIhaTitle"],
        )
    )
    story.append(
        _industrial_ai_pdf_paragraph(
            f"Diagnostic ID: {diagnostic_run.get('id', '—')}<br/>"
            "Brief type: industrial_ai<br/>"
            f"Brief version: {payload.get('brief_version', 'v1')}<br/>"
            f"Submitted at: {payload.get('submitted_at', 'не указано')}",
            styles["AIhaBody"],
        )
    )
    story.append(Spacer(1, 8))

    for section_title, fields in INDUSTRIAL_AI_BRIEF_SECTIONS:
        story.append(_industrial_ai_pdf_paragraph(section_title, styles["AIhaHeading"]))

        data = [
            [
                _industrial_ai_pdf_paragraph("Поле", styles["AIhaBody"]),
                _industrial_ai_pdf_paragraph("Значение", styles["AIhaBody"]),
            ]
        ]

        for label, path in fields:
            value = _industrial_ai_format_value(_industrial_ai_get(payload, path))
            data.append(
                [
                    _industrial_ai_pdf_paragraph(label, styles["AIhaBody"]),
                    _industrial_ai_pdf_paragraph(value, styles["AIhaBody"]),
                ]
            )

        table = Table(data, colWidths=[68 * mm, 100 * mm], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEBFF")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111111")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8C8D0")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )

        story.append(table)
        story.append(Spacer(1, 10))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)

def _industrial_form_text(name: str) -> str:
    return request.form.get(name, "").strip()


def _industrial_form_list(name: str) -> list[str]:
    return [
        value.strip()
        for value in request.form.getlist(name)
        if value and value.strip()
    ]


def _build_industrial_ai_brief_payload() -> dict[str, Any]:
    return {
        "brief_type": "industrial_ai",
        "brief_version": "v1",
        "source": "web_form",
        "submitted_at": datetime.utcnow().isoformat(),
        "industrial_ai": {
            "case_type": {
                "primary_case_type": _industrial_form_text("primary_case_type"),
                "primary_case_type_other": _industrial_form_text("primary_case_type_other"),
                "business_goals": _industrial_form_list("business_goals"),
                "expected_effect": _industrial_form_text("expected_effect"),
            },
            "production_context": {
                "company_name": _industrial_form_text("company_name"),
                "production_site": _industrial_form_text("production_site"),
                "unit_or_line": _industrial_form_text("unit_or_line"),
                "process_owner": _industrial_form_text("process_owner"),
                "technical_contact": _industrial_form_text("technical_contact"),
                "data_contact": _industrial_form_text("data_contact"),
                "current_process_description": _industrial_form_text("current_process_description"),
                "process_start": _industrial_form_text("process_start"),
                "process_end": _industrial_form_text("process_end"),
                "in_scope": _industrial_form_text("in_scope"),
                "out_of_scope": _industrial_form_text("out_of_scope"),
            },
            "analysis_object": {
                "equipment": {
                    "equipment_type": _industrial_form_text("equipment_type"),
                    "equipment_count": _industrial_form_text("equipment_count"),
                    "critical_equipment": _industrial_form_text("critical_equipment"),
                    "operation_mode": _industrial_form_text("operation_mode"),
                    "main_failures": _industrial_form_text("main_failures"),
                    "average_downtime": _industrial_form_text("average_downtime"),
                    "downtime_hour_cost": _industrial_form_text("downtime_hour_cost"),
                },
                "lines_operations": {
                    "lines_count": _industrial_form_text("lines_count"),
                    "main_operations": _industrial_form_text("main_operations"),
                    "bottlenecks": _industrial_form_text("bottlenecks"),
                    "changeover_frequency": _industrial_form_text("changeover_frequency"),
                    "changeover_duration": _industrial_form_text("changeover_duration"),
                },
                "products_materials": {
                    "product_types": _industrial_form_text("product_types"),
                    "batch_tracking": _industrial_form_text("batch_tracking"),
                    "critical_materials": _industrial_form_text("critical_materials"),
                    "critical_suppliers": _industrial_form_text("critical_suppliers"),
                    "material_issues": _industrial_form_text("material_issues"),
                },
            },
            "data": {
                "data_sources": _industrial_form_list("data_sources"),
                "data_period": _industrial_form_text("data_period"),
                "data_frequency": _industrial_form_text("data_frequency"),
                "data_identifiers": _industrial_form_text("data_identifiers"),
                "data_timestamps": _industrial_form_text("data_timestamps"),
                "data_quality_issues": _industrial_form_list("data_quality_issues"),
            },
            "systems_integrations": {
                "erp_status": _industrial_form_text("erp_status"),
                "mes_status": _industrial_form_text("mes_status"),
                "scada_status": _industrial_form_text("scada_status"),
                "wms_status": _industrial_form_text("wms_status"),
                "excel_export_status": _industrial_form_text("excel_export_status"),
                "api_status": _industrial_form_text("api_status"),
                "critical_integrations": _industrial_form_text("critical_integrations"),
                "deferred_integrations": _industrial_form_text("deferred_integrations"),
            },
            "security": {
                "sensitive_data": _industrial_form_list("sensitive_data"),
                "nda_status": _industrial_form_text("nda_status"),
                "cloud_allowed": _industrial_form_text("cloud_allowed"),
                "data_transfer_allowed": _industrial_form_text("data_transfer_allowed"),
                "preferred_processing": _industrial_form_text("preferred_processing"),
                "anonymization_requirements": _industrial_form_text("anonymization_requirements"),
                "security_notes": _industrial_form_text("security_notes"),
            },
            "economics": {
                "loss_areas": _industrial_form_list("loss_areas"),
                "downtime_events_per_month": _industrial_form_text("downtime_events_per_month"),
                "defect_rate": _industrial_form_text("defect_rate"),
                "defect_cost_per_month": _industrial_form_text("defect_cost_per_month"),
                "inventory_value": _industrial_form_text("inventory_value"),
                "manual_labor_hours": _industrial_form_text("manual_labor_hours"),
                "monthly_loss_estimate": _industrial_form_text("monthly_loss_estimate"),
                "meaningful_effect": _industrial_form_text("meaningful_effect"),
            },
            "mvp": {
                "mvp_case": _industrial_form_text("mvp_case"),
                "mvp_reason": _industrial_form_text("mvp_reason"),
                "mvp_unit": _industrial_form_text("mvp_unit"),
                "mvp_data_period": _industrial_form_text("mvp_data_period"),
                "mvp_test_period": _industrial_form_text("mvp_test_period"),
                "human_control": _industrial_form_text("human_control"),
                "mvp_exclusions": _industrial_form_text("mvp_exclusions"),
                "mvp_success_criteria": _industrial_form_text("mvp_success_criteria"),
                "mvp_go_criterion": _industrial_form_text("mvp_go_criterion"),
                "mvp_no_go_criterion": _industrial_form_text("mvp_no_go_criterion"),
            },
            "additional_comments": _industrial_form_text("additional_comments"),
        },
    }


@diagnostic_bp.route("/industrial-ai/<token>", methods=["GET", "POST"])
def industrial_ai_brief(token: str):
    diagnostic_run = get_diagnostic_run_by_token(token)

    if diagnostic_run is None:
        return render_template(
            "consulting/diagnostic_input_pack_invalid.html",
            site_links=get_site_links(),
        ), 404

    if request.method == "POST":
        payload = _build_industrial_ai_brief_payload()

        # Жёсткая защита, чтобы Industrial AI Brief не сохранился как обычный input pack
        payload["brief_type"] = "industrial_ai"
        payload["brief_version"] = payload.get("brief_version") or "v1"
        payload["source"] = payload.get("source") or "web_form"

        saved_input_pack = upsert_active_input_pack(
            diagnostic_run_id=diagnostic_run["id"],
            brief_type="industrial_ai",
            payload=payload,
            source=payload["source"],
        )

        input_pack_id = _extract_input_pack_id(saved_input_pack)

        uploaded_files = request.files.getlist("attachments")
        _save_uploaded_files(
            files=uploaded_files,
            diagnostic_run_id=diagnostic_run["id"],
            input_pack_id=input_pack_id,
        )

        return redirect(
            url_for("diagnostic.industrial_ai_brief_submitted", token=token)
        )

    active_input_pack = get_active_input_pack(
        diagnostic_run_id=diagnostic_run["id"],
        brief_type="industrial_ai",
    )

    form_data = _load_raw_payload(active_input_pack)

    existing_attachments = []
    if active_input_pack is not None:
        existing_attachments = get_input_pack_attachments(
            int(active_input_pack["id"])
        )

    return render_template(
        "consulting/industrial_ai_brief.html",
        site_links=get_site_links(),
        diagnostic_run=diagnostic_run,
        diagnostic=diagnostic_run,
        token=token,
        form_data=form_data,
        edit_mode=active_input_pack is not None,
        input_pack_id=active_input_pack["id"] if active_input_pack else None,
        existing_attachments=existing_attachments,
        download_docx_url=url_for(
            "diagnostic.download_industrial_ai_brief_template",
            token=token,
            file_format="docx",
        ),
        download_pdf_url=url_for(
            "diagnostic.download_industrial_ai_brief_template",
            token=token,
            file_format="pdf",
        ),
    )


@diagnostic_bp.route("/industrial-ai/<token>/download/<file_format>")
def download_industrial_ai_brief_template(token: str, file_format: str):
    diagnostic_run = get_diagnostic_run_by_token(token)

    if diagnostic_run is None:
        abort(404)

    file_format = file_format.lower().strip()

    if file_format not in {"docx", "pdf"}:
        abort(404)

    output_path = _make_temp_output_path(f".{file_format}")

    try:
        if file_format == "docx":
            builder = globals().get("_build_industrial_ai_blank_docx")
            download_name = "AIha_Industrial_AI_Brief_Form_v1.docx"
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            builder = globals().get("_build_industrial_ai_blank_pdf")
            download_name = "AIha_Industrial_AI_Brief_Form_v1.pdf"
            mimetype = "application/pdf"

        if builder is None:
            output_path.unlink(missing_ok=True)
            return (
                "Генератор файла для Industrial AI Brief ещё не подключён. "
                "Добавьте функции _build_industrial_ai_blank_docx / "
                "_build_industrial_ai_blank_pdf.",
                501,
            )

        builder(output_path=output_path, diagnostic_run=diagnostic_run)

    except RuntimeError as exc:
        output_path.unlink(missing_ok=True)
        return str(exc), 500

    response = send_file(
        output_path,
        as_attachment=True,
        download_name=download_name,
        mimetype=mimetype,
    )
    response.call_on_close(lambda: output_path.unlink(missing_ok=True))
    return response


@diagnostic_bp.route("/industrial-ai/<token>/submitted")
def industrial_ai_brief_submitted(token: str):
    diagnostic_run = get_diagnostic_run_by_token(token)

    if diagnostic_run is None:
        abort(404)

    return render_template(
        "consulting/industrial_ai_brief_submitted.html",
        site_links=get_site_links(),
        diagnostic_run=diagnostic_run,
        diagnostic=diagnostic_run,
        token=token,
        input_url=url_for("diagnostic.industrial_ai_brief", token=token),
        download_docx_url=url_for(
            "diagnostic.download_submitted_industrial_ai_brief",
            token=token,
            file_format="docx",
        ),
        download_pdf_url=url_for(
            "diagnostic.download_submitted_industrial_ai_brief",
            token=token,
            file_format="pdf",
        ),
    )


@diagnostic_bp.route("/industrial-ai/<token>/submitted/download/<file_format>")
def download_submitted_industrial_ai_brief(token: str, file_format: str):
    diagnostic_run = get_diagnostic_run_by_token(token)

    if diagnostic_run is None:
        abort(404)

    file_format = file_format.lower().strip()

    if file_format not in {"docx", "pdf"}:
        abort(404)

    input_pack = get_active_input_pack(
        diagnostic_run["id"],
        "industrial_ai",
    )

    if input_pack is None:
        return "Заполненный Industrial AI Brief не найден", 404

    payload = _load_raw_payload(input_pack)

    if payload.get("brief_type") != "industrial_ai":
        return "Заполненный Industrial AI Brief не найден", 404

    output_path = _make_temp_output_path(f".{file_format}")

    try:
        if file_format == "docx":
            builder = globals().get("_build_submitted_industrial_ai_docx")
            download_name = f"AIha_Industrial_AI_Brief_Submitted_{diagnostic_run['id']}.docx"
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            builder = globals().get("_build_submitted_industrial_ai_pdf")
            download_name = f"AIha_Industrial_AI_Brief_Submitted_{diagnostic_run['id']}.pdf"
            mimetype = "application/pdf"

        if builder is None:
            output_path.unlink(missing_ok=True)
            return (
                "Генератор заполненного файла для Industrial AI Brief ещё не подключён. "
                "Добавьте функции _build_submitted_industrial_ai_docx / "
                "_build_submitted_industrial_ai_pdf.",
                501,
            )

        builder(
            output_path=output_path,
            diagnostic_run=diagnostic_run,
            payload=payload,
        )

    except RuntimeError as exc:
        output_path.unlink(missing_ok=True)
        return str(exc), 500

    response = send_file(
        output_path,
        as_attachment=True,
        download_name=download_name,
        mimetype=mimetype,
    )
    response.call_on_close(lambda: output_path.unlink(missing_ok=True))
    return response
